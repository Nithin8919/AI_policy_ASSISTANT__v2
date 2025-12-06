"""
File Handler for processing uploaded files.
Extracts text from PDF, TXT, and DOCX files for temporary query augmentation.
"""

import os
import tempfile
import logging
from pathlib import Path
from typing import Dict, Optional
from fastapi import UploadFile

logger = logging.getLogger(__name__)

# Import text extraction utilities
import sys
sys.path.insert(0, str(Path(__file__).parent.parent.parent / 'ingestion_v2'))

from ingestion_v2.extraction.extract_text import TextExtractor
from ingestion_v2.io.file_loader import load_text_file

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX files won't be supported.")


class FileHandler:
    """Handle uploaded files and extract text content."""
    
    # File size limits
    MAX_FILE_SIZE_MB = 10
    MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
    
    # Supported file types
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx'}
    
    def __init__(self):
        """Initialize file handler."""
        self.pdf_extractor = TextExtractor()
        self.temp_dir = tempfile.gettempdir()
    
    def validate_file(self, file: UploadFile) -> tuple[bool, Optional[str]]:
        """
        Validate uploaded file.
        
        Args:
            file: Uploaded file
            
        Returns:
            (is_valid, error_message)
        """
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type: {file_ext}. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
        
        # Check DOCX availability
        if file_ext == '.docx' and not DOCX_AVAILABLE:
            return False, "DOCX support not available. Please install python-docx."
        
        # Check file size (if available)
        if hasattr(file, 'size') and file.size:
            if file.size > self.MAX_FILE_SIZE_BYTES:
                return False, f"File too large: {file.size / (1024*1024):.1f}MB. Max: {self.MAX_FILE_SIZE_MB}MB"
        
        return True, None
    
    async def _extract_with_gemini(self, file_path: str, mime_type: str = "application/pdf") -> Dict:
        """
        Extract text using Gemini 1.5 Flash (Cloud OCR).
        Used as fallback for scanned documents/images.
        """
        try:
            import google.generativeai as genai
            import time
            
            api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
            if not api_key:
                return {
                    "text": "",
                    "word_count": 0,
                    "success": False,
                    "error": "Gemini API key not found for OCR fallback"
                }

            genai.configure(api_key=api_key)
            
            # 1. Upload file
            logger.info(f"📤 Uploading {Path(file_path).name} to Gemini for OCR...")
            uploaded_file = genai.upload_file(file_path, mime_type=mime_type)
            
            # 2. Wait for processing (usually fast for small files)
            while uploaded_file.state.name == "PROCESSING":
                time.sleep(1)
                uploaded_file = genai.get_file(uploaded_file.name)
                
            if uploaded_file.state.name == "FAILED":
                return {
                    "text": "",
                    "word_count": 0,
                    "success": False,
                    "error": "Gemini file processing failed"
                }

            # 3. Generate content (Extract text)
            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(
                [uploaded_file, "Extract all text from this document verbatim. Preserve structure where possible."],
                generation_config={"temperature": 0.0}
            )
            
            text = response.text
            word_count = len(text.split())
            
            logger.info(f"✨ Gemini OCR extracted {word_count} words")
            
            # Cleanup (optional, but good practice if we uploaded it)
            # genai.delete_file(uploaded_file.name) 
            
            return {
                "text": text,
                "word_count": word_count,
                "success": True,
                "method": "gemini_ocr"
            }

        except Exception as e:
            logger.error(f"Gemini OCR failed: {e}")
            return {
                "text": "",
                "word_count": 0,
                "success": False,
                "error": f"Gemini OCR error: {str(e)}"
            }

    async def process_file(self, file: UploadFile) -> Dict:
        """
        Process uploaded file and extract text.
        
        Args:
            file: Uploaded file
            
        Returns:
            Dictionary with:
                - text: Extracted text
                - filename: Original filename
                - file_type: File extension
                - word_count: Number of words
                - success: Boolean
                - error: Error message if failed
        """
        # Validate file
        is_valid, error_msg = self.validate_file(file)
        if not is_valid:
            return {
                'text': '',
                'filename': file.filename,
                'file_type': Path(file.filename).suffix.lower(),
                'word_count': 0,
                'success': False,
                'error': error_msg
            }
        
        temp_path = None
        try:
            # Save to temporary file
            temp_path = await self._save_temp_file(file)
            
            # Extract text based on file type
            file_ext = Path(file.filename).suffix.lower()
            
            result = None
            
            if file_ext == '.pdf':
                result = self._extract_pdf(temp_path)
            elif file_ext == '.txt':
                result = self._extract_txt(temp_path)
            elif file_ext == '.docx':
                result = self._extract_docx(temp_path)
            else:
                result = {
                    'text': '',
                    'success': False,
                    'error': f'Unsupported file type: {file_ext}'
                }
            
            # CHECK FOR SCANNED PDF / POOR EXTRACTION
            # If word count is low (< 50) and it's a PDF, try Gemini OCR
            if file_ext == '.pdf' and (not result['success'] or result.get('word_count', 0) < 50):
                logger.info(f"⚠️ Low text extraction ({result.get('word_count', 0)} words). Attempting Gemini OCR fallback...")
                ocr_result = await self._extract_with_gemini(temp_path, mime_type="application/pdf")
                
                # If OCR worked better, use it
                if ocr_result['success'] and ocr_result['word_count'] > result.get('word_count', 0):
                    result = ocr_result
            
            # Add metadata
            result['filename'] = file.filename
            result['file_type'] = file_ext
            
            logger.info(f"Processed {file.filename}: {result.get('word_count', 0)} words")
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {file.filename}: {e}")
            return {
                'text': '',
                'filename': file.filename,
                'file_type': Path(file.filename).suffix.lower(),
                'word_count': 0,
                'success': False,
                'error': str(e)
            }
        finally:
            # Clean up temporary file
            if temp_path and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                    logger.debug(f"Cleaned up temp file: {temp_path}")
                except Exception as e:
                    logger.warning(f"Failed to clean up temp file: {e}")
    
    async def _save_temp_file(self, file: UploadFile) -> str:
        """Save uploaded file to temporary location."""
        # Create unique temp filename
        file_ext = Path(file.filename).suffix
        temp_fd, temp_path = tempfile.mkstemp(suffix=file_ext, dir=self.temp_dir)
        
        try:
            # Read file content
            content = await file.read()
            
            # Write to temp file
            with os.fdopen(temp_fd, 'wb') as f:
                f.write(content)
            
            # Reset file pointer for potential re-reading
            await file.seek(0)
            
            return temp_path
        except Exception as e:
            # Clean up on error
            if os.path.exists(temp_path):
                os.remove(temp_path)
            raise e
    
    def _extract_pdf(self, file_path: str) -> Dict:
        """Extract text from PDF file."""
        try:
            result = self.pdf_extractor.extract(Path(file_path))
            
            if result['success']:
                return {
                    'text': result['text'],
                    'word_count': result['word_count'],
                    'page_count': result.get('page_count', 0),
                    'success': True,
                    'method': result.get('method', 'unknown')
                }
            else:
                return {
                    'text': '',
                    'word_count': 0,
                    'success': False,
                    'error': 'PDF extraction failed'
                }
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return {
                'text': '',
                'word_count': 0,
                'success': False,
                'error': str(e)
            }
    
    def _extract_txt(self, file_path: str) -> Dict:
        """Extract text from TXT file."""
        try:
            text = load_text_file(Path(file_path))
            
            if text:
                word_count = len(text.split())
                return {
                    'text': text,
                    'word_count': word_count,
                    'success': True,
                    'method': 'direct_read'
                }
            else:
                return {
                    'text': '',
                    'word_count': 0,
                    'success': False,
                    'error': 'Failed to read text file'
                }
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return {
                'text': '',
                'word_count': 0,
                'success': False,
                'error': str(e)
            }
    
    def _extract_docx(self, file_path: str) -> Dict:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            return {
                'text': '',
                'word_count': 0,
                'success': False,
                'error': 'python-docx not installed'
            }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine all text
            all_text = '\n\n'.join(paragraphs)
            if table_texts:
                all_text += '\n\nTables:\n' + '\n'.join(table_texts)
            
            word_count = len(all_text.split())
            
            return {
                'text': all_text,
                'word_count': word_count,
                'success': True,
                'method': 'python-docx'
            }
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return {
                'text': '',
                'word_count': 0,
                'success': False,
                'error': str(e)
            }
